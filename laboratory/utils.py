import os
import json
import logging
import subprocess
from django.http import JsonResponse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from django.views.decorators.csrf import csrf_exempt
from reception.models import Patient, AnalysisResult, Analysis
from django.conf import settings
from django.utils import timezone
from datetime import timedelta


def convert_docx_to_pdf(docx_path: str) -> str:
    output_dir = os.path.dirname(docx_path)
    subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path], check=True)

    return docx_path.replace(".docx", ".pdf")


def create_analysis_docx(patient, analysis, results_list, output_path,
                         header_image_path=None, analysis_title="Analiz"):
    doc = Document()

    # ===== PAGE SETUP =====
    section = doc.sections[0]
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    style.font.size = Pt(11)

    # ===== HEADER =====
    if header_image_path and os.path.exists(header_image_path):
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(3)
        table.columns[1].width = Inches(3.5)

        img_p = table.cell(0, 0).paragraphs[0]
        img_p.add_run().add_picture(header_image_path, width=Inches(2.5))
        img_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        text_p = table.cell(0, 1).paragraphs[0]
        text_p.add_run(
            "MANZIL: QARSHI SHAHAR KAT - MFY, NASAF KO' CHASI, 31-UY\n"
            "TEL: (75) 223-47-47\n"
            "MOBIL: (97) 070-47-47 ; (97) 310-21-01"
        ).font.size = Pt(10)

    # ===== PATIENT INFO =====
    info = doc.add_table(rows=4, cols=2, style='Table Grid')
    info.cell(0, 0).text = "Bemor I.F.O"
    info.cell(0, 1).text = f"{patient.name or ''} {patient.last_name or ''} {patient.middle_name or ''}"
    info.cell(1, 0).text = "Tug'ilgan sanasi"
    info.cell(1, 1).text = patient.birth_date.strftime("%Y-%m-%d") if patient.birth_date else ""
    info.cell(2, 0).text = "Telefon raqami"
    info.cell(2, 1).text = patient.phone_number or ""
    info.cell(3, 0).text = "Tekshiruv sanasi"
    info.cell(3, 1).text = analysis.created_at.strftime("%Y-%m-%d %H:%M")

    subtitle = doc.add_paragraph()
    subtitle.add_run(f"{analysis_title} : {patient.id}").italic = True
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ===== RESULTS TABLE =====
    if not results_list:
        p = doc.add_paragraph("Natija hali kiritilmagan")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    else:
        table = doc.add_table(rows=1, cols=3, style='Table Grid')
        table.autofit = False
        table.columns[0].width = Inches(4.0)
        table.columns[1].width = Inches(2.0)
        table.columns[2].width = Inches(4.0)

        hdr = table.rows[0].cells
        hdr[0].text = "Tahlil nomi"
        hdr[1].text = "Tahlil natijasi"
        hdr[2].text = "Norma"

        for item in results_list:
            title = item['title']
            value = item['value']
            norma_lines = [n.strip() for n in (item['norma'] or "").splitlines() if n.strip()]
            if not norma_lines:
                norma_lines = ["-"]

            # first row
            first_row = table.add_row().cells
            first_row[0].text = title
            first_row[1].text = value
            first_row[2].text = norma_lines[0]

            start_row = len(table.rows) - 1

            # remaining norma rows
            for n in norma_lines[1:]:
                r = table.add_row().cells
                r[2].text = n

            end_row = len(table.rows) - 1

            # merge title & result
            if end_row > start_row:
                table.cell(start_row, 0).merge(table.cell(end_row, 0))
                table.cell(start_row, 1).merge(table.cell(end_row, 1))

    # ===== SIGN =====
    dept = analysis.department_types.department
    user = dept.user_set.first()

    sign = doc.add_table(rows=1, cols=2)
    sign.autofit = False
    sign.columns[0].width = Inches(3)
    sign.columns[1].width = Inches(3)

    sign.cell(0, 0).paragraphs[0].add_run(
        "Врач лаборант:  _____________________").bold = False

    sign.cell(0, 1).paragraphs[0].add_run(
        user.full_name if user else ""
    )

    stamp = os.path.join(settings.MEDIA_ROOT, "images", "pechat.jpg")
    if os.path.exists(stamp):
        p = doc.add_paragraph()
        p.add_run().add_picture(stamp, width=Inches(3))
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.save(output_path)


logger = logging.getLogger(__name__)


@csrf_exempt
def export_analysis_by_phone(request):
    logger.error("=== EXPORT STARTED ===")

    if request.method != "POST":
        return JsonResponse({"error": "Only POST"}, status=405)

    try:
        body = json.loads(request.body)
        patient_id = body.get("patient_id")
        logger.error(f"Request body: patient_id={patient_id}")
    except Exception:
        logger.exception("JSON parse error")
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    if not patient_id:
        return JsonResponse({"error": "patient_id required"}, status=400)

    try:
        patient = Patient.objects.get(id=int(patient_id))
    except Patient.DoesNotExist:
        logger.error("Patient not found")
        return JsonResponse({"error": "Patient not found"}, status=404)

    logger.error(f"Patient FOUND: id={patient.id}, name={patient.name}")

    one_month_ago = timezone.now() - timedelta(days=30)
    analyses = Analysis.objects.filter(patient=patient, created_at__gte=one_month_ago).order_by('-created_at')
    logger.error(f"Total analyses found: {analyses.count()}")

    if not analyses.exists():
        return JsonResponse({"error": "No analysis found"}, status=404)

    export_dir = os.path.join(settings.MEDIA_ROOT, "temp_exports")
    os.makedirs(export_dir, exist_ok=True)

    header_image_path = os.path.join(settings.MEDIA_ROOT, "images", "logo.jpg")
    if not os.path.exists(header_image_path):
        header_image_path = None

    files_created = []

    for analysis in analyses:
        logger.error(f"--- PROCESSING ANALYSIS id={analysis.id}")

        full_name = " ".join(filter(None, [
            patient.name,
            patient.last_name,
            getattr(patient, "middle_name", "")
        ])).strip()

        if not full_name:
            full_name = f"patient_{patient.id}"

        base = f"{full_name}_{analysis.id}"
        docx_path = os.path.join(export_dir, f"{base}.docx")

        results_list = []
        seen_titles = set()

        qs = (
            AnalysisResult.objects
            .filter(
                patient=patient,
                result__department_types=analysis.department_types,
                analysis_result__isnull=False
            )
            .exclude(analysis_result__exact="")
            .select_related("result")
            .order_by("-created_at")
        )

        for ar in qs:
            if not ar.result:
                continue

            title = ar.result.title
            if title in seen_titles:
                continue

            seen_titles.add(title)

            norma_text = (ar.result.norma or "").strip()

            results_list.append({
                "title": title,
                "value": ar.analysis_result,
                "norma": norma_text
            })

        logger.error(f"FINAL results_list count: {len(results_list)}")

        create_analysis_docx(
            patient=patient,
            analysis=analysis,
            results_list=results_list,
            output_path=docx_path,
            header_image_path=header_image_path,
            analysis_title=f"{analysis.department_types.title} natijasi"
        )

        pdf_path = convert_docx_to_pdf(docx_path)

        if os.path.exists(docx_path):
            os.remove(docx_path)

        pdf_url = request.build_absolute_uri(
            f"{settings.MEDIA_URL}temp_exports/{os.path.basename(pdf_path)}"
        )

        files_created.append({
            "filename": os.path.basename(pdf_path),
            "url": pdf_url
        })

    logger.error("=== EXPORT FINISHED ===")

    return JsonResponse({"files": files_created})
