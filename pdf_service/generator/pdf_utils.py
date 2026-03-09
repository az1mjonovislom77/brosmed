import os
import subprocess
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn


def convert_docx_to_pdf(docx_path: str):
    output_dir = os.path.dirname(docx_path)
    subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", "--outdir",
                    output_dir, docx_path], check=True)

    return docx_path.replace(".docx", ".pdf")


def create_analysis_docx(patient, analysis, results_list, output_path, header_image_path=None,
                         analysis_title="Analiz"):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(11)

    if header_image_path is None:
        header_image_path = "/home/brosmed/laboratory/logo.jpg"

    if os.path.exists(header_image_path):
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.columns[0].width = Inches(3.4)
        table.columns[1].width = Inches(3.1)
        cell_img = table.cell(0, 0)
        p_img = cell_img.paragraphs[0]
        run_img = p_img.add_run()
        run_img.add_picture(header_image_path, width=Inches(2.5))
        p_img.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        cell_img.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_text = table.cell(0, 1)
        cell_text.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p_text = cell_text.paragraphs[0]

        text = (
            "MANZIL: QARSHI SHAHAR KAT - MFY,\n"
            "NASAF KO'CHASI, 31-UY\n"
            "TEL: (75) 223-47-47\n"
            "MOBIL: (97) 070-47-47 ; (97) 310-21-01"
        )

        run_text = p_text.add_run(text)
        run_text.font.size = Pt(10)
        p_text.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = "Table Grid"
    info_table.autofit = True
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.cell(0, 0).text = "Bemor I.F.O"
    info_table.cell(0, 1).text = (
        f"{patient.get('name', '')} "
        f"{patient.get('last_name', '')} "
        f"{patient.get('middle_name', '')}"
    )
    info_table.cell(1, 0).text = "Tugilgan sanasi"
    info_table.cell(1, 1).text = patient.get("birth_date", "")
    info_table.cell(2, 0).text = "Telefon raqami"
    info_table.cell(2, 1).text = patient.get("phone_number", "")
    info_table.cell(3, 0).text = "Tekshiruv sanasi"
    info_table.cell(3, 1).text = analysis.get("created_at", "")

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(f"{analysis_title} : {patient.get('id')}\n")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(13)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    valid_results = []

    for item in results_list:
        raw_value = item.get("value")

        if raw_value is None:
            continue

        value_str = str(raw_value).strip()

        if not value_str or value_str.lower() in ["none", "null", ""]:
            continue

        title = item.get("title", "").strip()
        norma = item.get("norma", "")
        norma_str = str(norma).strip() if norma is not None else ""

        if not norma_str:
            norma_str = "-"

        norma_str = (
            norma_str.replace("\n", " | ")
            .replace("Мужчины:", "М:")
            .replace("Женщины:", "Ж:")
            .replace("Жен:", "Ж:")
            .replace("Муж:", "М:")
        )

        valid_results.append({"title": title, "value": value_str, "norma": norma_str, })

    if not valid_results:
        p = doc.add_paragraph("Natija hali kiritilmagan")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.runs[0]
        run.italic = True
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

    else:
        table = doc.add_table(rows=1, cols=3, style="Table Grid")
        table.autofit = False
        table.columns[0].width = Inches(3.9)
        table.columns[1].width = Inches(1.9)
        table.columns[2].width = Inches(1.9)
        hdr = table.rows[0].cells
        hdr[0].text = "Название анализа"
        hdr[1].text = "Результат анализа"
        hdr[2].text = "Норма"

        for item in reversed(valid_results):
            row = table.add_row().cells
            row[0].text = item["title"]
            row[1].text = item["value"]
            row[2].text = item["norma"]

            for cell in row:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10)

    sign_table = doc.add_table(rows=1, cols=2)
    sign_table.autofit = False
    sign_table.columns[0].width = Inches(3)
    sign_table.columns[1].width = Inches(3)
    left_cell = sign_table.cell(0, 0)
    left_p = left_cell.paragraphs[0]
    left_p.add_run("Врач лаборант: _____________________")
    left_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    right_cell = sign_table.cell(0, 1)
    right_p = right_cell.paragraphs[0]
    right_p.add_run(analysis.get("doctor", ""))
    right_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.save(output_path)
